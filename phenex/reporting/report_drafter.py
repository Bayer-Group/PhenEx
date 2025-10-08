import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import io
import base64
import os
import json
import re
from typing import List, Optional, Dict, Any, Union
from datetime import datetime, date
from pathlib import Path
import logging

from phenex.reporting.reporter import Reporter
from phenex.reporting.waterfall import Waterfall
from phenex.reporting.table1 import Table1
from phenex.reporting.table2 import Table2
from phenex.util import create_logger

logger = create_logger(__name__)


# Legacy ReportLab code removed - now using markdown-pdf directly


# Optional imports for document generation

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn

    DOCX_AVAILABLE = True
except ImportError:
    logger.warning(
        "python-docx not available. Word document generation will not work. Install with: pip install python-docx"
    )
    DOCX_AVAILABLE = False

try:
    from openai import AzureOpenAI, OpenAI

    OPENAI_AVAILABLE = True
except ImportError:
    logger.warning(
        "OpenAI library not available. AI-generated text will fall back to rules-based. Install with: pip install openai"
    )
    OPENAI_AVAILABLE = False

try:
    from markdown_pdf import MarkdownPdf, Section

    MARKDOWN_PDF_AVAILABLE = True
except ImportError:
    logger.warning(
        "markdown-pdf not available. PDF generation will not work. Install with: pip install markdown-pdf"
    )
    MARKDOWN_PDF_AVAILABLE = False


class ReportDrafter(Reporter):
    """
    The ReportDrafter creates comprehensive final study reports including:
    - Cohort definition description (entry, inclusion, exclusion criteria)
    - Data analysis description and date ranges
    - Waterfall table showing patient attrition
    - Study variables (characteristics and outcomes)
    - Table 1 (baseline characteristics)
    - Table 2 (outcomes analysis)
    - AI-generated descriptive text and figure captions (when OpenAI is available)

    The report can be exported to PDF or Word format.

    Parameters:
        use_ai: Whether to use OpenAI for generating descriptive text (default: True if API keys available)
        ai_model: OpenAI model to use for text generation (default: "gpt-4o-mini")
        include_plots: Whether to include plots in the report (default: True)
        plot_dpi: DPI for plots (default: 300)
        title: Report title (if None, will be generated from cohort name)
        author: Report author(s)
        institution: Institution name
        date_range_start: Start date for data analysis (if None, inferred from cohort)
        date_range_end: End date for data analysis (if None, inferred from cohort)
    """

    def __init__(
        self,
        use_ai: bool = True,
        ai_model: str = "gpt-4o-mini",
        include_plots: bool = True,
        plot_dpi: int = 300,
        title: Optional[str] = None,
        author: Optional[str] = None,
        institution: Optional[str] = None,
        date_range_start: Optional[Union[str, date]] = None,
        date_range_end: Optional[Union[str, date]] = None,
        decimal_places: int = 1,
        pretty_display: bool = True,
    ):
        super().__init__(decimal_places=decimal_places, pretty_display=pretty_display)

        self.use_ai = use_ai and OPENAI_AVAILABLE and self._check_openai_config()
        self.ai_model = ai_model
        self.include_plots = include_plots
        self.plot_dpi = plot_dpi
        self.title = title
        self.author = author
        self.institution = institution
        self.date_range_start = date_range_start
        self.date_range_end = date_range_end

        # Initialize OpenAI client if available
        self.ai_client = None
        if self.use_ai:
            self._initialize_ai_client()

        # Report sections storage
        self.report_sections = {}
        self.figures = {}

    def _check_openai_config(self) -> bool:
        """Check if OpenAI configuration is available in environment variables."""
        logger.debug("🔍 Checking OpenAI configuration in environment variables...")

        # Check for Azure OpenAI or standard OpenAI configuration
        azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
        api_version = os.getenv("OPENAI_API_VERSION")
        openai_api_key = os.getenv("OPENAI_API_KEY")

        logger.debug(f"Environment check results:")
        logger.debug(
            f"  AZURE_OPENAI_ENDPOINT: {'✅ Set' if azure_endpoint else '❌ Not set'}"
        )
        logger.debug(
            f"  AZURE_OPENAI_API_KEY: {'✅ Set' if azure_api_key else '❌ Not set'}"
        )
        logger.debug(
            f"  OPENAI_API_VERSION: {'✅ Set' if api_version else '❌ Not set'} ({api_version})"
        )
        logger.debug(
            f"  OPENAI_API_KEY: {'✅ Set' if openai_api_key else '❌ Not set'}"
        )

        has_azure = bool(azure_endpoint and azure_api_key)
        has_openai = bool(openai_api_key)

        logger.debug(f"Configuration status: Azure={has_azure}, OpenAI={has_openai}")

        return has_azure or has_openai

    def _initialize_ai_client(self):
        """Initialize OpenAI client (Azure or standard)."""
        logger.info("Attempting to initialize AI client for text generation")

        # Try to load environment variables from common locations
        env_paths = [
            Path.cwd() / ".env",
            Path.cwd() / "app" / ".env",
            Path(__file__).parent.parent.parent / "app" / ".env",
        ]

        logger.debug(f"Attempting to load environment from common locations...")
        for env_path in env_paths:
            if env_path.exists():
                logger.info(f"📁 Found .env file at: {env_path}")
                try:
                    from dotenv import load_dotenv

                    load_dotenv(env_path)
                    logger.info(f"✅ Successfully loaded environment from: {env_path}")
                    break
                except Exception as e:
                    logger.warning(f"Failed to load .env from {env_path}: {e}")
            else:
                logger.debug(f"No .env file found at: {env_path}")

        try:
            # Check for Azure OpenAI configuration
            azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
            api_version = os.getenv("OPENAI_API_VERSION", "2024-02-15-preview")

            logger.debug(f"Current environment state after loading:")
            logger.debug(
                f"  AZURE_OPENAI_ENDPOINT: {azure_endpoint[:50] + '...' if azure_endpoint else 'Not set'}"
            )
            logger.debug(
                f"  AZURE_OPENAI_API_KEY: {'*' * min(20, len(azure_api_key)) if azure_api_key else 'Not set'}"
            )
            logger.debug(f"  OPENAI_API_VERSION: {api_version}")
            logger.debug(f"  Selected AI Model: {self.ai_model}")

            if azure_endpoint and azure_api_key:
                logger.info(
                    f"Found Azure OpenAI credentials - endpoint: {azure_endpoint[:50]}..., API version: {api_version}"
                )
                # Clean up API version (remove quotes if present)
                if api_version.startswith('"') and api_version.endswith('"'):
                    api_version = api_version.strip('"')

                # Initialize Azure OpenAI client with minimal parameters
                logger.info("Initializing Azure OpenAI client...")

                # First try with normal SSL verification
                self.ai_client = AzureOpenAI(
                    azure_endpoint=azure_endpoint.strip(),
                    api_key=azure_api_key.strip(),
                    api_version=api_version.strip(),
                )
                logger.debug("✅ Azure OpenAI client initialized with SSL verification")

                # Keep the model name as specified in constructor (gpt-4o-mini by default)
                self._is_azure = True
                logger.info(
                    f"Azure OpenAI client initialized successfully with API version: {api_version}"
                )

                # Test the connection with a simple call
                try:
                    logger.info(
                        "Testing Azure OpenAI connection with minimal API call..."
                    )
                    logger.debug(
                        f"Test call details: model={self.ai_model}, endpoint={azure_endpoint}"
                    )
                    logger.debug(f"API version: {api_version}")

                    test_response = self.ai_client.chat.completions.create(
                        model=self.ai_model,
                        messages=[{"role": "user", "content": "Test connection"}],
                        max_tokens=5,
                    )
                    logger.info(
                        "✅ Azure OpenAI connection test successful - AI functionality is fully operational"
                    )
                    logger.debug(
                        f"Test response received: {len(test_response.choices[0].message.content)} chars"
                    )

                except Exception as test_error:
                    logger.error(
                        f"❌ Azure OpenAI connection test failed: {test_error}"
                    )
                    logger.debug(f"Error type: {type(test_error).__name__}")
                    logger.debug(f"Full error details: {str(test_error)}")

                    # Provide specific debugging based on error type
                    error_str = str(test_error).lower()
                    if "connection" in error_str:
                        logger.error("🌐 Connection Error Debugging:")
                        logger.error(f"   • Endpoint: {azure_endpoint}")
                        logger.error(f"   • Check network connectivity")
                        logger.error(f"   • Verify firewall/proxy settings")
                        logger.error(f"   • Confirm endpoint URL is correct")
                    elif "authentication" in error_str or "unauthorized" in error_str:
                        logger.error("🔑 Authentication Error Debugging:")
                        logger.error(
                            f"   • API key length: {len(azure_api_key) if azure_api_key else 0} chars"
                        )
                        logger.error(f"   • Check API key validity")
                        logger.error(f"   • Verify key permissions")
                    elif "model" in error_str:
                        logger.error("🤖 Model Error Debugging:")
                        logger.error(f"   • Requested model: {self.ai_model}")
                        logger.error(f"   • Check model deployment name")
                        logger.error(
                            f"   • Verify model availability in your Azure instance"
                        )

                    raise test_error

            else:
                # Fall back to standard OpenAI
                openai_api_key = os.getenv("OPENAI_API_KEY")
                if openai_api_key:
                    logger.info(
                        "Found standard OpenAI API key, initializing standard OpenAI client..."
                    )
                    self.ai_client = OpenAI(api_key=openai_api_key.strip())
                    # Keep the model name as specified in constructor
                    self._is_azure = False
                    logger.info("✅ Standard OpenAI client initialized successfully")
                else:
                    logger.warning("No OpenAI API keys found in environment variables")
                    raise Exception("No valid API keys found")

        except Exception as e:
            logger.warning(f"Failed to initialize OpenAI client: {e}")
            # Log more details for debugging
            if "proxies" in str(e):
                logger.info(
                    "Note: 'proxies' error may be due to library version or environment configuration"
                )
            self.ai_client = None
            self.use_ai = False
            self._is_azure = False

    def _generate_ai_text(
        self,
        prompt: str,
        max_tokens: int = 16384,
    ) -> str:
        """Generate text using AI or fallback to rules-based generation."""
        if not self.use_ai:
            logger.debug("AI disabled, using fallback text generation")
            return self._fallback_text_generation(prompt, None)

        logger.info(
            f"🤖 Making AI API call for text generation (max_tokens: {max_tokens})..."
        )

        # Inject global context automatically from class variable
        global_context = getattr(self, "_global_context", "")
        full_prompt = f"{global_context}\n\n{prompt}"

        logger.debug(f"AI prompt preview: {prompt[:100]}...")

        try:
            # Try to use the OpenAI client
            response = self.ai_client.chat.completions.create(
                model=self.ai_model,
                messages=[
                    {
                        "role": "system",
                        "content": self._get_global_ai_system_instructions(),
                    },
                    {"role": "user", "content": full_prompt},
                ],
                max_tokens=max_tokens,
                temperature=0,
            )
            generated_text = response.choices[0].message.content.strip()
            logger.info(
                f"✅ AI text generation successful (generated {len(generated_text)} characters)"
            )
            logger.debug(f"AI response preview: {generated_text[:100]}...")
            return generated_text
        except Exception as e:
            logger.warning(
                f"❌ AI text generation failed, falling back to rules-based: {e}"
            )
            return self._fallback_text_generation(prompt, None)

    def _get_global_ai_system_instructions(self) -> str:
        """Get global system instructions for all AI text generation calls."""
        return """You are a professional medical researcher and biostatistician writing for a high-impact scientific publication.

WRITING STYLE REQUIREMENTS:
- Keep answers concise yet insightful and informative
- Use professional medical research language suitable for peer-reviewed journals
- Provide clinical context and significance for all findings
- Be precise with statistical interpretations
- Include relevant clinical implications
- Use appropriate medical terminology consistently
- Write in active voice where appropriate
- Ensure content is publication-ready

CONTENT STANDARDS:
- All statistical claims must be clinically meaningful
- Include appropriate caveats and limitations where relevant  
- Focus on actionable clinical insights
- Maintain scientific objectivity and accuracy
- Reference established clinical guidelines and norms when relevant
- Ensure content flows logically and cohesively

FORMATTING:
- Use clean markdown formatting
- Structure content with clear headings and sections
- Use bullet points for lists where appropriate
- Ensure proper medical/scientific citation style"""

    def _build_global_ai_context(self, cohort) -> str:
        """
        Build comprehensive global context that is automatically injected into every AI call.
        This ensures all AI responses have complete study awareness and consistency.
        """
        context_parts = []

        # Study Overview
        context_parts.append(
            f"""=== COMPREHENSIVE STUDY CONTEXT ===
STUDY TITLE: {getattr(self, 'title', 'Medical Research Study')}
COHORT NAME: {getattr(cohort, 'name', 'Study Cohort')}
COHORT DESCRIPTION: {getattr(cohort, 'description', 'Not available')}
STUDY TYPE: Comprehensive medical research study analyzing patient outcomes and characteristics"""
        )

        # Cohort Information
        n_entry_patients = (
            cohort.entry_criterion.table.select("PERSON_ID")
            .distinct()
            .count()
            .execute()
        )
        n_index_patients = cohort.table.select("PERSON_ID").distinct().count().execute()
        context_parts.append(
            f"ENTRY COHORT SIZE (ENTRY CRITERION ONLY): {n_entry_patients} patients"
        )
        context_parts.append(
            f"FINAL COHORT SIZE (ALL INEX CRITERIA APPLIED): {n_index_patients} patients"
        )

        # Entry criteria
        context_parts.append(f"ENTRY CRITERION: {cohort.entry_criterion.name}")
        context_parts.append(
            f"\n\t{json.dumps(cohort.entry_criterion.to_dict(), indent=4)}"
        )

        # Inclusions
        if hasattr(cohort, "inclusions") and cohort.inclusions:
            context_parts.append(
                f"\nINCLUSION CRITERIA ({len(cohort.inclusions)} criteria):"
            )
            for i, inclusion in enumerate(cohort.inclusions, 1):
                name = getattr(inclusion, "name", f"Inclusion {i}")
                context_parts.append(f"  {i}. {name}")
                context_parts.append(f"\n\t{json.dumps(inclusion.to_dict(), indent=4)}")

        # Exclusions
        if hasattr(cohort, "exclusions") and cohort.exclusions:
            context_parts.append(
                f"\nEXCLUSION CRITERIA ({len(cohort.exclusions)} criteria):"
            )
            for i, exclusion in enumerate(cohort.exclusions, 1):
                name = getattr(exclusion, "name", f"Exclusion {i}")
                context_parts.append(f"  {i}. {name}")
                context_parts.append(f"\n\t{json.dumps(exclusion.to_dict(), indent=4)}")

        # Characteristics
        if hasattr(cohort, "characteristics") and cohort.characteristics:
            context_parts.append(
                f"\nBASELINE CHARACTERISTICS ({len(cohort.characteristics)} variables):"
            )
            for i, char in enumerate(cohort.characteristics, 1):
                name = getattr(char, "name", f"Characteristic {i}")
                context_parts.append(f"  {i}. {name}")
                context_parts.append(f"\n\t{json.dumps(char.to_dict(), indent=4)}")

        # Outcomes
        if hasattr(cohort, "outcomes") and cohort.outcomes:
            context_parts.append(
                f"\nOUTCOME MEASURES ({len(cohort.outcomes)} variables):"
            )
            for i, outcome in enumerate(cohort.outcomes, 1):
                name = getattr(outcome, "name", f"Outcome {i}")
                context_parts.append(f"  {i}. {name}")
                context_parts.append(f"\n\t{json.dumps(outcome.to_dict(), indent=4)}")

        # Report generation metadata
        if hasattr(self, "author") and self.author:
            context_parts.append(f"\nREPORT AUTHOR: {self.author}")
        if hasattr(self, "institution") and self.institution:
            context_parts.append(f"INSTITUTION: {self.institution}")

        context_parts.append("\n=== END GLOBAL CONTEXT ===")

        return "\n".join(context_parts)

    def _generate_ai_image_caption(self, image_base64: str, context: str) -> str:
        """Generate caption for image using AI text generation (no vision API needed)."""
        if not self.use_ai or not self.ai_client:
            logger.debug(
                "AI disabled or client unavailable, using simple figure caption"
            )
            return f"Figure: {context}"

        logger.info(f"🖼️ Generating AI figure caption using text model...")
        logger.debug(f"Image context: {context}")

        try:
            # Use text-based generation instead of vision API since gpt-4-vision-preview is not available
            prompt = f"""
            Generate a professional medical research figure caption for a plot with this context: {context}
            
            The caption should be:
            - Professional and suitable for a medical research publication
            - Clear and descriptive
            - Include relevant clinical interpretation
            - Follow standard academic figure caption format
            
            Start with "Figure X:" and provide a comprehensive description.
            """

            response = self.ai_client.chat.completions.create(
                model=self.ai_model,  # Use the same model as other text generation
                messages=[
                    {
                        "role": "system",
                        "content": "You are a professional medical researcher writing figure captions for a scientific publication.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
            )
            generated_caption = response.choices[0].message.content.strip()
            logger.info(
                f"✅ AI figure caption generation successful (generated {len(generated_caption)} characters)"
            )
            return generated_caption
        except Exception as e:
            logger.warning(
                f"❌ AI figure caption generation failed, using fallback: {e}"
            )
            return f"Figure: {context}"

    def _fallback_text_generation(self, prompt: str, cohort=None) -> str:
        """Fallback text generation using rules-based approach."""
        if "cohort definition" in prompt.lower():
            return (
                self._create_specific_cohort_description(cohort)
                if cohort
                else "This cohort was defined using entry criteria, inclusion criteria, and exclusion criteria as specified in the methods section."
            )
        elif "data analysis" in prompt.lower():
            return "Data analysis was performed using the PhenEx framework, applying the specified phenotype definitions and filters."
        elif "baseline characteristics" in prompt.lower():
            return "Baseline characteristics were calculated at the index date for all patients meeting the inclusion and exclusion criteria."
        elif "outcomes" in prompt.lower():
            return "Outcomes were evaluated for all patients in the final cohort during the follow-up period."
        else:
            return "Details are provided in the accompanying tables and figures."

    def _format_cohort_name(self, name: str) -> str:
        """Format cohort name from snake_case to proper Title Case."""
        if not name:
            return name

        # Convert snake_case or camelCase to Title Case
        import re

        # Handle snake_case (e.g., study_tutorial_cohort -> Study Tutorial Cohort)
        formatted = re.sub(r"_", " ", name)

        # Handle camelCase (e.g., studyTutorialCohort -> Study Tutorial Cohort)
        formatted = re.sub(r"([a-z])([A-Z])", r"\1 \2", formatted)

        # Convert to title case
        formatted = formatted.title()

        return formatted

    def _create_executive_summary(self) -> str:
        """Generate AI-powered executive summary in journal abstract style."""
        logger.info("Generating AI executive summary")

        prompt = """TASK: Write a professional medical journal-style executive summary/abstract for this study.
        
        ABSTRACT STRUCTURE REQUIREMENTS:
        - **Objective:** What was studied and why
        - **Methods:** Brief description of study design, population, and criteria
        - **Results:** Key findings from baseline characteristics and outcomes (use realistic clinical interpretations)
        - **Conclusions:** Clinical implications and significance
        
        SPECIFIC REQUIREMENTS:
        - Medical journal abstract format (150-250 words)
        - Focus on clinical significance and real-world implications
        - Use realistic medical findings appropriate for the study population
        - Include key statistical insights where clinically relevant
        
        Write a complete executive summary that reads like a published medical research abstract."""

        return self._generate_ai_text(prompt)

    def _create_cohort_description(self, cohort) -> str:
        """Generate cohort definition description."""
        logger.info(f"Generating cohort description for: {cohort.name}")

        # Format the cohort name properly
        formatted_cohort_name = self._format_cohort_name(cohort.name)

        prompt = f"""
        Write a professional description of this medical research cohort definition using clean markdown formatting:
        
        Cohort Name: {formatted_cohort_name}
        Cohort Description: {cohort.description or 'Not provided'}
        
        Entry Criterion: {cohort.entry_criterion.to_dict()}
        
        Inclusion Criteria:
        {chr(10).join([f"- {inc.display_name if hasattr(inc, 'display_name') else inc.name}" for inc in (cohort.inclusions or [])])}
        
        Exclusion Criteria:
        {chr(10).join([f"- {exc.display_name if hasattr(exc, 'display_name') else exc.name}" for exc in (cohort.exclusions or [])])}
        
        Please write a professional medical research description with:
        - Brief introduction paragraph about the study population
        - **Entry Criterion:** section with rationale
        - **Inclusion Criteria:** section with bullet points (use * for bullets)
        - **Exclusion Criteria:** section with bullet points (use * for bullets)
        - Clinical rationale for each criterion
        
        Use clean markdown formatting with proper line breaks between sections.
        """
        return self._generate_ai_text(prompt)

    def _create_specific_cohort_description(self, cohort) -> str:
        """Create specific cohort description with actual criteria listed."""
        if not cohort:
            return "This cohort was defined using entry criteria, inclusion criteria, and exclusion criteria as specified in the methods section."

        description_parts = []

        # Add cohort description if available
        if hasattr(cohort, "description") and cohort.description:
            description_parts.append(f"Study Population: {cohort.description}")

        # Entry criterion
        entry_name = (
            cohort.entry_criterion.display_name
            if hasattr(cohort.entry_criterion, "display_name")
            else cohort.entry_criterion.name
        )
        description_parts.append(f"Entry Criterion: Patients with {entry_name}.")

        # Inclusion criteria
        if cohort.inclusions:
            inclusion_list = []
            for inc in cohort.inclusions:
                inc_name = (
                    inc.display_name if hasattr(inc, "display_name") else inc.name
                )
                inclusion_list.append(inc_name)
            if inclusion_list:
                description_parts.append(
                    f"Inclusion Criteria: {', '.join(inclusion_list)}."
                )

        # Exclusion criteria
        if cohort.exclusions:
            exclusion_list = []
            for exc in cohort.exclusions:
                exc_name = (
                    exc.display_name if hasattr(exc, "display_name") else exc.name
                )
                exclusion_list.append(exc_name)
            if exclusion_list:
                description_parts.append(
                    f"Exclusion Criteria: Patients were excluded if they had {', '.join(exclusion_list)}."
                )

        return " ".join(description_parts)

    def _create_data_analysis_description(self, cohort) -> str:
        """Generate data analysis description."""
        logger.info("Generating data analysis description")

        n_patients = (
            cohort.index_table.filter(cohort.index_table.BOOLEAN == True)
            .select("PERSON_ID")
            .distinct()
            .count()
            .execute()
        )

        prompt = f"""
        Write a description of the data analysis for this medical research study:
        
        - Final cohort size: {n_patients} patients
        - Analysis period: {self.date_range_start or 'Study period'} to {self.date_range_end or 'End of follow-up'}
        - Characteristics analyzed: {len(cohort.characteristics or [])} baseline characteristics
        - Outcomes analyzed: {len(cohort.outcomes or [])} outcome measures
        
        Describe the analytical approach, patient population, and study period.
        """
        return self._generate_ai_text(prompt)

    def _create_variables_description(self, cohort) -> str:
        """Generate description of study variables."""
        logger.info("Generating study variables description")

        characteristics = cohort.characteristics or []
        outcomes = cohort.outcomes or []

        char_names = [
            c.display_name if hasattr(c, "display_name") else c.name
            for c in characteristics
        ]
        outcome_names = [
            o.display_name if hasattr(o, "display_name") else o.name for o in outcomes
        ]

        prompt = f"""
        Write a professional description of the study variables for this medical research study.
        
        Baseline Characteristics ({len(characteristics)}):
        {chr(10).join([f"- {name}" for name in char_names])}
        
        Outcome Variables ({len(outcomes)}):
        {chr(10).join([f"- {name}" for name in outcome_names])}
        
        REQUIREMENTS:
        - Use numbered lists for each variable (1. Variable Name: Description)
        - Group baseline characteristics separately from outcome variables
        - Explain the clinical relevance of each variable
        - Include measurement methods where appropriate
        
        Write a comprehensive study variables section.
        """
        return self._generate_ai_text(prompt)

    def _build_comprehensive_study_context(
        self, cohort=None, additional_context=""
    ) -> str:
        """
        Build comprehensive study context for AI prompts to ensure consistency and accuracy.
        This context should be included at the start of all AI generation prompts.
        """
        context_parts = []

        # Study Overview
        context_parts.append(
            f"""
=== COMPREHENSIVE STUDY CONTEXT ===

STUDY TITLE: {getattr(self, 'title', 'Medical Research Study')}
COHORT NAME: {getattr(self, 'cohort_name', 'Study Cohort')}

STUDY DESIGN: This is a comprehensive medical research study analyzing patient outcomes and characteristics.
"""
        )

        # Cohort Information
        if cohort:
            try:
                # Get patient count
                n_patients = (
                    cohort.index_table.filter(cohort.index_table.BOOLEAN == True)
                    .select("PERSON_ID")
                    .distinct()
                    .count()
                    .execute()
                )
                context_parts.append(f"FINAL COHORT SIZE: {n_patients} patients")
            except:
                context_parts.append("FINAL COHORT SIZE: [To be determined]")

            # Entry criteria
            if hasattr(cohort, "name"):
                context_parts.append(f"PRIMARY COHORT DEFINITION: {cohort.name}")

            # Inclusions
            if cohort.inclusions:
                context_parts.append(
                    f"\nINCLUSION CRITERIA ({len(cohort.inclusions)} criteria):"
                )
                for i, inclusion in enumerate(cohort.inclusions, 1):
                    name = (
                        inclusion.display_name
                        if hasattr(inclusion, "display_name")
                        else inclusion.name
                    )
                    context_parts.append(f"  {i}. {name}")

            # Exclusions
            if cohort.exclusions:
                context_parts.append(
                    f"\nEXCLUSION CRITERIA ({len(cohort.exclusions)} criteria):"
                )
                for i, exclusion in enumerate(cohort.exclusions, 1):
                    name = (
                        exclusion.display_name
                        if hasattr(exclusion, "display_name")
                        else exclusion.name
                    )
                    context_parts.append(f"  {i}. {name}")

            # Characteristics
            if cohort.characteristics:
                context_parts.append(
                    f"\nBASELINE CHARACTERISTICS ({len(cohort.characteristics)} variables):"
                )
                for i, char in enumerate(cohort.characteristics, 1):
                    name = (
                        char.display_name
                        if hasattr(char, "display_name")
                        else char.name
                    )
                    context_parts.append(f"  {i}. {name}")

            # Outcomes
            if cohort.outcomes:
                context_parts.append(
                    f"\nOUTCOME MEASURES ({len(cohort.outcomes)} variables):"
                )
                for i, outcome in enumerate(cohort.outcomes, 1):
                    name = (
                        outcome.display_name
                        if hasattr(outcome, "display_name")
                        else outcome.name
                    )
                    context_parts.append(f"  {i}. {name}")

        # Existing report sections for consistency
        if hasattr(self, "report_sections") and self.report_sections:
            existing_sections = []
            for k in self.report_sections.keys():
                value = self.report_sections.get(k)
                # Handle DataFrames and other objects properly
                if value is not None:
                    if hasattr(value, "empty"):  # DataFrame
                        if not value.empty:
                            existing_sections.append(k)
                    elif isinstance(value, str) and value.strip():  # String
                        existing_sections.append(k)
                    elif value:  # Other truthy values
                        existing_sections.append(k)
            if existing_sections:
                context_parts.append(
                    f"\nEXISTING REPORT SECTIONS: {', '.join(existing_sections)}"
                )

        # Additional context
        if additional_context:
            context_parts.append(f"\nADDITIONAL CONTEXT: {additional_context}")

        context_parts.append("\n=== END STUDY CONTEXT ===\n")

        return "\n".join(context_parts)

    def _generate_waterfall_commentary(self, waterfall_df):
        """Generate AI commentary for waterfall table."""
        logger.info("Generating waterfall table commentary")

        if waterfall_df is None or waterfall_df.empty:
            return "No waterfall data available for analysis."

        # Extract key statistics
        initial_n = waterfall_df.iloc[0]["N"] if len(waterfall_df) > 0 else "Unknown"
        final_n = (
            waterfall_df.iloc[-1]["Remaining"] if len(waterfall_df) > 0 else "Unknown"
        )
        inclusion_steps = waterfall_df[waterfall_df["Type"] == "inclusion"]
        exclusion_steps = waterfall_df[waterfall_df["Type"] == "exclusion"]

        prompt = f"""
        Analyze this patient attrition waterfall table and write a professional clinical commentary.
        
        WATERFALL DATA:
        Initial patient pool: {initial_n}
        Final cohort size: {final_n}
        Number of inclusion criteria: {len(inclusion_steps)}
        Number of exclusion criteria: {len(exclusion_steps)}
        
        Detailed attrition steps:
        {waterfall_df[['Type', 'Name', 'N', 'Remaining']].to_string()}
        
        Focus on:
        - Clinical interpretation of patient selection process
        - Analysis of attrition rates and their implications
        - Assessment of study representativeness and generalizability
        - Discussion of potential selection bias considerations
        """

        return self._generate_ai_text(prompt)

    def _generate_table1_commentary(self, table1_df):
        """Generate AI commentary for Table 1 baseline characteristics."""
        logger.info("Generating Table 1 commentary")

        if table1_df is None or table1_df.empty:
            return "No baseline characteristics data available for analysis."

        prompt = f"""
        Analyze this Table 1 baseline characteristics and write a professional clinical commentary.
        
        BASELINE CHARACTERISTICS DATA:
        {table1_df.to_string()}
        
        Focus on:
        - Clinical interpretation of baseline demographics and characteristics
        - Assessment of population representativeness
        - Clinical implications for study outcomes
        - Comparison to relevant population norms where appropriate
        - Risk factor assessment and clinical significance
        """

        return self._generate_ai_text(prompt)

    def _generate_table2_commentary(self, table2_df):
        """Generate AI commentary for Table 2 outcomes."""
        logger.info("Generating Table 2 commentary")

        if table2_df is None or table2_df.empty:
            return "No outcomes data available for analysis."

        prompt = f"""
        Analyze this Table 2 outcomes summary and write a professional clinical commentary.
        
        OUTCOMES DATA:
        {table2_df.to_string()}
        
        Focus on:
        - Clinical interpretation of outcome results
        - Assessment of key findings and their significance
        - Clinical implications for patient care and clinical practice
        - Risk assessment and prognostic implications
        - Comparison to published literature where appropriate
        """

        return self._generate_ai_text(prompt)

    def _plot_to_base64(self, fig) -> str:
        """Convert matplotlib figure to base64 string."""
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=self.plot_dpi, bbox_inches="tight")
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        buffer.close()
        return image_base64

    def _create_waterfall_plot(self, waterfall_df: pd.DataFrame) -> tuple:
        """Create waterfall plot and return figure and base64 string."""
        fig, ax = plt.subplots(figsize=(10, 6))

        # Create waterfall chart
        y_pos = range(len(waterfall_df))
        remaining = waterfall_df["Remaining"].values

        bars = ax.barh(y_pos, remaining, color="steelblue", alpha=0.7)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(
            [f"{row['Type']}: {row['Name']}" for _, row in waterfall_df.iterrows()]
        )
        ax.set_xlabel("Number of Patients")
        ax.set_title("Patient Attrition (Waterfall Chart)")

        # Add value labels on bars
        for i, (bar, value) in enumerate(zip(bars, remaining)):
            ax.text(
                bar.get_width() + max(remaining) * 0.01,
                bar.get_y() + bar.get_height() / 2,
                f"{int(value):,}",
                va="center",
                fontsize=9,
            )

        ax.grid(axis="x", alpha=0.3)
        plt.tight_layout()

        image_base64 = self._plot_to_base64(fig)
        return fig, image_base64

    def execute(self, cohort) -> Dict[str, Any]:
        """Execute the report generation."""
        logger.info(f"Generating comprehensive report for cohort: {cohort.name}")

        # Ensure cohort is executed
        if cohort.index_table is None:
            logger.error("Cohort not yet executed. Run cohort execution first.")

        # Generate title if not provided
        if not self.title:
            self.title = f"Study Report: {cohort.name}"

        # STEP 1: Build the global AI context once at the beginning and set as class variable
        logger.info("Building global AI context...")
        self._global_context = self._build_global_ai_context(cohort)

        # STEP 2: Generate data tables first (for context enhancement)
        logger.info("=== PHASE 1: Generating Data Tables ===")

        # Generate Waterfall Table
        logger.info("Generating waterfall table...")
        waterfall_reporter = Waterfall(
            decimal_places=self.decimal_places, pretty_display=self.pretty_display
        )
        waterfall_df = waterfall_reporter.execute(cohort)
        self.report_sections["waterfall_table"] = waterfall_df

        # Generate Table 1 (Baseline Characteristics)
        if cohort.characteristics:
            logger.info("Generating Table 1 (baseline characteristics)...")
            table1_reporter = Table1(
                decimal_places=self.decimal_places, pretty_display=True
            )  # Enable pretty display for proper formatting
            try:
                table1_df = table1_reporter.execute(cohort)
                self.report_sections["table1"] = table1_df
                logger.info(
                    f"Table1 generated successfully with {len(table1_df)} rows and columns: {list(table1_df.columns)}"
                )
            except Exception as e:
                logger.error(f"FATAL: Table1 generation failed: {e}")
                logger.error(f"Error type: {type(e).__name__}")

                # The Table1 reporter is a core component and should work with properly structured cohorts
                # If it's failing, the issue is likely with our mock data structure
                raise RuntimeError(
                    f"Table1 reporter failed. This indicates the cohort characteristics are not properly structured for the Table1 reporter. Original error: {e}"
                )
        else:
            logger.info("No characteristics defined. Skipping Table 1.")
            self.report_sections["table1"] = pd.DataFrame()

        # 8. Generate Table 2 (Outcomes) if outcomes exist
        if cohort.outcomes:
            logger.info("Generating Table 2 (outcomes)...")

            # Initialize Table2 reporter with the exposure phenotype
            table2_reporter = Table2(
                time_points=[365],  # 1 year follow-up
                decimal_places=self.decimal_places,
                pretty_display=True,
            )

            try:
                table2_df = table2_reporter.execute(cohort)
                self.report_sections["table2"] = table2_df
                logger.info(
                    f"Table2 generated successfully with {len(table2_df)} rows and columns: {list(table2_df.columns)}"
                )
            except Exception as e:
                logger.error(f"FATAL: Table2 generation failed: {e}")
                logger.error(f"Error type: {type(e).__name__}")

                # Table2 reporter is a core component and should work with properly structured cohorts
                raise RuntimeError(
                    f"Table2 reporter failed. This indicates the cohort structure is not compatible with the Table2 reporter. Original error: {e}"
                )
        else:
            logger.info("No outcomes defined. Skipping Table 2.")
            self.report_sections["table2"] = pd.DataFrame()

        # Generate summary statistics
        n_patients = (
            cohort.index_table.filter(cohort.index_table.BOOLEAN == True)
            .select("PERSON_ID")
            .distinct()
            .count()
            .execute()
        )
        self.report_sections["summary_stats"] = {
            "total_patients": n_patients,
            "n_characteristics": len(cohort.characteristics or []),
            "n_outcomes": len(cohort.outcomes or []),
            "n_inclusions": len(cohort.inclusions or []),
            "n_exclusions": len(cohort.exclusions or []),
        }

        # STEP 3: Generate AI-Powered Content (using class variable for global context)
        logger.info("=== PHASE 2: Generating AI-Powered Content ===")

        # Generate AI text sections using global context class variable
        logger.info("Generating AI executive summary...")
        self.report_sections["executive_summary"] = self._create_executive_summary()

        logger.info("Generating cohort definition description...")
        self.report_sections["cohort_definition"] = self._create_cohort_description(
            cohort
        )

        logger.info("Generating data analysis description...")
        self.report_sections["data_analysis"] = self._create_data_analysis_description(
            cohort
        )

        logger.info("Generating study variables description...")
        self.report_sections["study_variables"] = self._create_variables_description(
            cohort
        )

        # Generate AI commentary for tables and figures
        if self.ai_client:
            logger.info("Generating AI commentary for waterfall table...")
            self.report_sections["waterfall_commentary"] = (
                self._generate_waterfall_commentary(
                    self.report_sections.get("waterfall_table")
                )
            )

            logger.info("Generating AI commentary for Table 1...")
            self.report_sections["table1_commentary"] = (
                self._generate_table1_commentary(self.report_sections.get("table1"))
            )

            logger.info("Generating AI commentary for Table 2...")
            self.report_sections["table2_commentary"] = (
                self._generate_table2_commentary(self.report_sections.get("table2"))
            )

        # Generate plots if requested (with AI captions that now have full context)
        if self.include_plots and not waterfall_df.empty:
            logger.info("Generating waterfall plot...")
            fig, img_b64 = self._create_waterfall_plot(waterfall_df)
            self.figures["waterfall"] = {
                "figure": fig,
                "base64": img_b64,
                "caption": self._generate_ai_image_caption(
                    img_b64,
                    "Patient attrition waterfall chart showing how inclusion and exclusion criteria affected the final cohort size",
                ),
            }

        logger.info("Report generation completed successfully")
        return self.report_sections

    def to_pdf(self, filename: str, output_dir: str = ".") -> str:
        """
        Generate a PDF report using markdown-pdf for clean, professional formatting.

        Args:
            filename: Name of the PDF file to create
            output_dir: Directory to save the PDF in

        Returns:
            Path to the generated PDF file
        """
        if not MARKDOWN_PDF_AVAILABLE:
            raise ImportError(
                "markdown-pdf is required for PDF generation. Install with: pip install markdown-pdf"
            )

        if not self.report_sections:
            raise ValueError("No report data available. Call execute() first.")

        output_path = Path(output_dir) / filename
        if not output_path.suffix:
            output_path = output_path.with_suffix(".pdf")

        logger.info(f"Generating PDF report using markdown-pdf: {output_path}")

        # Create the PDF document
        pdf = MarkdownPdf(toc_level=2, optimize=True)

        # Set document metadata
        pdf.meta["title"] = self.title or "Study Report"
        pdf.meta["author"] = self.author or "PhenEx Report Generator"
        if self.institution:
            pdf.meta["subject"] = f"Medical Research Report - {self.institution}"
        pdf.meta["keywords"] = "medical research, cohort study, phenex"

        # Get global CSS for all sections
        css = self._get_global_css()

        # Add title section (not in TOC)
        title_section = self._build_title_section()
        pdf.add_section(Section(title_section, toc=True), user_css=css)

        # Add each major section separately for proper TOC
        section_number = 1

        # 1. Cohort Definition
        if "cohort_definition" in self.report_sections:
            cohort_section = f"# {section_number}. Cohort Definition\n\n"
            cohort_section += self.report_sections["cohort_definition"]
            pdf.add_section(Section(cohort_section, toc=True), user_css=css)
            section_number += 1

        # 2. Data Analysis
        if "data_analysis" in self.report_sections:
            analysis_section = f"# {section_number}. Data Analysis\n\n"
            analysis_section += self.report_sections["data_analysis"]
            pdf.add_section(Section(analysis_section, toc=True), user_css=css)
            section_number += 1

        # 3. Study Variables
        if "study_variables" in self.report_sections:
            variables_section = f"# {section_number}. Study Variables\n\n"
            variables_section += self.report_sections["study_variables"]
            pdf.add_section(Section(variables_section, toc=True), user_css=css)
            section_number += 1

        # 4. Patient Attrition (Waterfall Table)
        if (
            "waterfall_table" in self.report_sections
            and not self.report_sections["waterfall_table"].empty
        ):
            waterfall_section = f"# {section_number}. Patient Attrition\n\n"
            waterfall_df = self.report_sections["waterfall_table"]
            waterfall_section += self._dataframe_to_markdown_table(waterfall_df)
            waterfall_section += "\n\n"
            if "waterfall_commentary" in self.report_sections:
                waterfall_section += "## Clinical Commentary\n\n"
                waterfall_section += self.report_sections["waterfall_commentary"]
            pdf.add_section(Section(waterfall_section, toc=True), user_css=css)
            section_number += 1

        # 5. Baseline Characteristics (Table 1)
        if (
            "table1" in self.report_sections
            and not self.report_sections["table1"].empty
        ):
            table1_section = f"# {section_number}. Baseline Characteristics\n\n"
            table1_df = self.report_sections["table1"]
            table1_section += self._dataframe_to_markdown_table(table1_df)
            table1_section += "\n\n"
            if "table1_commentary" in self.report_sections:
                table1_section += "## Clinical Commentary\n\n"
                table1_section += self.report_sections["table1_commentary"]
            pdf.add_section(Section(table1_section, toc=True), user_css=css)
            section_number += 1

        # 6. Outcomes Summary (Table 2)
        if (
            "table2" in self.report_sections
            and not self.report_sections["table2"].empty
        ):
            table2_section = f"# {section_number}. Outcomes Summary\n\n"
            table2_df = self.report_sections["table2"]
            table2_section += self._dataframe_to_markdown_table(table2_df)
            table2_section += "\n\n"
            if "table2_commentary" in self.report_sections:
                table2_section += "## Clinical Commentary\n\n"
                table2_section += self.report_sections["table2_commentary"]
            pdf.add_section(Section(table2_section, toc=True), user_css=css)
            section_number += 1

        # Save the PDF
        pdf.save(str(output_path))
        logger.info(f"PDF report generated: {output_path}")
        return str(output_path)

    def _build_title_section(self) -> str:
        """Build the title section for the PDF report."""
        title_section = ""
        if self.title:
            title_section += f"# {self.title}\n\n"
        if self.author:
            title_section += f"**Author:** {self.author}\n\n"
        if self.institution:
            title_section += f"**Institution:** {self.institution}\n\n"
        if hasattr(self, "date") and self.date:
            title_section += f"**Date:** {self.date}\n\n"
        else:
            title_section += (
                f"**Report Generated:** {datetime.now().strftime('%B %d, %Y')}\n\n"
            )

        # Add the AI-generated executive summary
        title_section += "## Executive Summary\n\n"
        if "executive_summary" in self.report_sections:
            title_section += self.report_sections["executive_summary"]
        else:
            # Fallback to basic summary if AI summary not available
            stats = self.report_sections.get("summary_stats", {})
            title_section += f"""This report presents the analysis of {stats.get('total_patients', 'N/A')} patients in the study cohort. 
The analysis includes {stats.get('n_characteristics', 0)} baseline characteristics and {stats.get('n_outcomes', 0)} outcome measures.
Cohort definition involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria."""

        title_section += "\n\n"
        if title_section:
            title_section += "---\n\n"  # Horizontal rule separator
        return title_section

    def _build_complete_report(self) -> str:
        """Build the complete report with numbered sections in a single document."""
        report_md = f"# {self.title}\n\n"

        if self.author:
            report_md += f"**Author:** {self.author}\n\n"
        if self.institution:
            report_md += f"**Institution:** {self.institution}\n\n"

        report_md += f"**Report Generated:** {datetime.now().strftime('%B %d, %Y')}\n\n"

        # Executive Summary - AI Generated
        report_md += "## Executive Summary\n\n"
        if (
            hasattr(self, "report_sections")
            and "executive_summary" in self.report_sections
        ):
            report_md += self.report_sections["executive_summary"]
        else:
            # Fallback to basic summary if AI summary not available
            stats = self.report_sections.get("summary_stats", {})
            report_md += f"""This report presents the analysis of {stats.get('total_patients', 'N/A')} patients in the study cohort. 
The analysis includes {stats.get('n_characteristics', 0)} baseline characteristics and {stats.get('n_outcomes', 0)} outcome measures.
Cohort definition involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria."""
        report_md += "\n\n"

        section_number = 1

        # 1. Cohort Definition
        if "cohort_definition" in self.report_sections:
            report_md += f"## {section_number}. Cohort Definition\n\n"
            report_md += self.report_sections["cohort_definition"]
            report_md += "\n\n"
            section_number += 1

        # 2. Data Analysis
        if "data_analysis" in self.report_sections:
            report_md += f"## {section_number}. Data Analysis\n\n"
            report_md += self.report_sections["data_analysis"]
            report_md += "\n\n"
            section_number += 1

        # 3. Study Variables
        if "study_variables" in self.report_sections:
            report_md += f"## {section_number}. Study Variables\n\n"
            report_md += self.report_sections["study_variables"]
            report_md += "\n\n"
            section_number += 1

        # 4. Patient Attrition (Waterfall Table)
        if (
            "waterfall_table" in self.report_sections
            and not self.report_sections["waterfall_table"].empty
        ):
            report_md += f"## {section_number}. Patient Attrition\n\n"

            # Convert DataFrame to markdown table
            waterfall_df = self.report_sections["waterfall_table"]
            report_md += self._dataframe_to_markdown_table(waterfall_df)
            report_md += "\n\n"

            # Add commentary if available
            if "waterfall_commentary" in self.report_sections:
                report_md += "### Clinical Commentary\n\n"
                report_md += self.report_sections["waterfall_commentary"]
                report_md += "\n\n"
            section_number += 1

        # 5. Baseline Characteristics (Table 1)
        if (
            "table1" in self.report_sections
            and not self.report_sections["table1"].empty
        ):
            report_md += f"## {section_number}. Baseline Characteristics\n\n"

            # Convert DataFrame to markdown table
            table1_df = self.report_sections["table1"]
            report_md += self._dataframe_to_markdown_table(table1_df)
            report_md += "\n\n"

            # Add commentary if available
            if "table1_commentary" in self.report_sections:
                report_md += "### Clinical Commentary\n\n"
                report_md += self.report_sections["table1_commentary"]
                report_md += "\n\n"
            section_number += 1

        # 6. Outcomes Summary (Table 2)
        if (
            "table2" in self.report_sections
            and not self.report_sections["table2"].empty
        ):
            report_md += f"## {section_number}. Outcomes Summary\n\n"

            # Convert DataFrame to markdown table
            table2_df = self.report_sections["table2"]
            report_md += self._dataframe_to_markdown_table(table2_df)
            report_md += "\n\n"

            # Add commentary if available
            if "table2_commentary" in self.report_sections:
                report_md += "### Clinical Commentary\n\n"
                report_md += self.report_sections["table2_commentary"]
                report_md += "\n\n"
            section_number += 1

        return report_md

    def _build_waterfall_section(self) -> str:
        """Build the waterfall table section in markdown."""
        section_md = "# Patient Attrition (Waterfall Table)\n\n"

        # Convert DataFrame to markdown table
        waterfall_df = self.report_sections["waterfall_table"]
        section_md += self._dataframe_to_markdown_table(waterfall_df)
        section_md += "\n\n"

        # Add commentary if available
        if "waterfall_commentary" in self.report_sections:
            section_md += "## Clinical Commentary\n\n"
            section_md += self.report_sections["waterfall_commentary"]
            section_md += "\n\n"

        return section_md

    def _build_table1_section(self) -> str:
        """Build the Table 1 section in markdown."""
        section_md = "# Baseline Characteristics (Table 1)\n\n"

        # Convert DataFrame to markdown table
        table1_df = self.report_sections["table1"]
        section_md += self._dataframe_to_markdown_table(table1_df)
        section_md += "\n\n"

        # Add commentary if available
        if "table1_commentary" in self.report_sections:
            section_md += "## Clinical Commentary\n\n"
            section_md += self.report_sections["table1_commentary"]
            section_md += "\n\n"

        return section_md

    def _build_table2_section(self) -> str:
        """Build the Table 2 section in markdown."""
        section_md = "# Outcomes Summary (Table 2)\n\n"

        # Convert DataFrame to markdown table
        table2_df = self.report_sections["table2"]
        section_md += self._dataframe_to_markdown_table(table2_df)
        section_md += "\n\n"

        # Add commentary if available
        if "table2_commentary" in self.report_sections:
            section_md += "## Clinical Commentary\n\n"
            section_md += self.report_sections["table2_commentary"]
            section_md += "\n\n"

        return section_md

    def _dataframe_to_markdown_table(self, df) -> str:
        """Convert DataFrame to markdown table format."""
        if df is None or df.empty:
            return "_No data available_\n"

        # Build header
        headers = "|" + "|".join(str(col) for col in df.columns) + "|\n"
        separator = "|" + "|".join("---" for _ in df.columns) + "|\n"

        # Build rows
        rows = ""
        for _, row in df.iterrows():
            row_data = "|" + "|".join(str(val) for val in row) + "|\n"
            rows += row_data

        return headers + separator + rows

    def _get_global_css(self) -> str:
        """Get comprehensive CSS styling for all tables in the report."""
        return """
        /* Global document styling */
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333;
            background: none;
        }
        
        /* Header styling */
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Times New Roman', serif;
            color: #333;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            font-family: 'Times New Roman', serif;
            font-size: 11pt;
            background: none;
        }
        
        th {
            font-weight: bold;
            padding: 10px 8px;
            text-align: left;
            border: 1px solid #333;
            color: #333;
        }
        
        td {
            padding: 8px;
            border: 1px solid #333;
            text-align: left;
            color: #333;
        }
        
        table tbody tr {
            border: none;
        }
        
        table thead tr {
            border: none;
        }
            
        /* Center-aligned tables (for waterfall and outcomes) */
        table.center-align td,
        table.center-align th {
            text-align: center;
        }
        
        /* Remove any background shadows or effects */
        p, div, section, article {
            background: none;
        }
        
        /* Ensure text is readable */
        * {
            color: #333 !important;
        }
        """

    def to_word(self, filename: str, output_dir: str = ".") -> str:
        """Generate Word document report."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word document generation. Install with: pip install python-docx"
            )

        if not self.report_sections:
            raise ValueError("No report data available. Call execute() first.")

        output_path = Path(output_dir) / filename
        if not output_path.suffix:
            output_path = output_path.with_suffix(".docx")

        logger.info(f"Generating Word document: {output_path}")

        # Create Word document
        doc = Document()

        # Title
        title = doc.add_heading(self.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Author: {self.author or 'Not specified'}")
        doc.add_paragraph(f"Institution: {self.institution or 'Not specified'}")
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_page_break()

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        stats = self.report_sections.get("summary_stats", {})
        summary_text = f"""
        This report presents the analysis of {stats.get('total_patients', 'N/A')} patients in the study cohort. 
        The analysis includes {stats.get('n_characteristics', 0)} baseline characteristics and {stats.get('n_outcomes', 0)} outcome measures.
        Cohort definition involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria.
        """
        doc.add_paragraph(summary_text)

        # Cohort Definition
        doc.add_heading("1. Cohort Definition", level=1)
        cohort_def = self.report_sections.get(
            "cohort_definition", "No description available."
        )
        doc.add_paragraph(cohort_def)

        # Data Analysis
        doc.add_heading("2. Data Analysis", level=1)
        data_analysis = self.report_sections.get(
            "data_analysis", "No description available."
        )
        doc.add_paragraph(data_analysis)

        # Study Variables
        doc.add_heading("3. Study Variables", level=1)
        study_vars = self.report_sections.get(
            "study_variables", "No description available."
        )
        doc.add_paragraph(study_vars)

        # Waterfall Table
        waterfall_df = self.report_sections.get("waterfall_table")
        if waterfall_df is not None and not waterfall_df.empty:
            doc.add_heading("4. Patient Attrition (Waterfall Table)", level=1)

            # Add table to Word document
            table = doc.add_table(rows=1, cols=len(waterfall_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(waterfall_df.columns):
                hdr_cells[i].text = str(col)

            # Data rows
            for _, row in waterfall_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Add waterfall plot if available
            if "waterfall" in self.figures:
                doc.add_paragraph()
                doc.add_paragraph("Figure 1: Patient Attrition Waterfall")

                # Save plot temporarily for inclusion
                temp_plot_path = Path(output_dir) / "temp_waterfall.png"
                # Ensure directory exists
                temp_plot_path.parent.mkdir(parents=True, exist_ok=True)
                self.figures["waterfall"]["figure"].savefig(
                    temp_plot_path, dpi=self.plot_dpi, bbox_inches="tight"
                )
                doc.add_picture(str(temp_plot_path), width=Inches(6))
                doc.add_paragraph(self.figures["waterfall"]["caption"])

        # Table 1
        table1_df = self.report_sections.get("table1")
        if table1_df is not None and not table1_df.empty:
            doc.add_heading("5. Baseline Characteristics (Table 1)", level=1)

            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Characteristic"
            hdr_cells[1].text = "N"
            hdr_cells[2].text = "%"

            # Data rows
            for idx, row in table1_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(row.get("N", ""))
                row_cells[2].text = str(row.get("%", ""))

        # Table 2 (Outcomes)
        table2_df = self.report_sections.get("table2")
        if table2_df is not None and not table2_df.empty:
            doc.add_heading("6. Outcomes Summary (Table 2)", level=1)

            # Create table
            table = doc.add_table(rows=1, cols=len(table2_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(table2_df.columns):
                hdr_cells[i].text = str(col)

            # Data rows
            for _, row in table2_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

        # Save document
        doc.save(str(output_path))

        # Clean up temporary plot file if it exists
        temp_plot_path = Path(output_dir) / "temp_waterfall.png"
        if temp_plot_path.exists():
            try:
                temp_plot_path.unlink()
            except:
                pass

        logger.info(f"Word document generated: {output_path}")
        return str(output_path)

    def get_report_summary(self) -> Dict[str, Any]:
        """Get a summary of the generated report."""
        if not self.report_sections:
            return {"error": "No report data available. Call execute() first."}

        summary = {
            "title": self.title,
            "author": self.author,
            "institution": self.institution,
            "generation_date": datetime.now().isoformat(),
            "ai_enabled": self.use_ai,
            "sections_generated": list(self.report_sections.keys()),
            "figures_generated": list(self.figures.keys()),
            "summary_statistics": self.report_sections.get("summary_stats", {}),
        }

        # Add table shapes
        for section_name, section_data in self.report_sections.items():
            if isinstance(section_data, pd.DataFrame):
                summary[f"{section_name}_shape"] = section_data.shape

        return summary
